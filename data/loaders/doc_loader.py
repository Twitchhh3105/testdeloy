"""Load .docx / .pdf / .txt files, preserving heading hierarchy + tables."""
from __future__ import annotations

import unicodedata
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterator


@dataclass
class DocBlock:
    """A single logical block within a document (heading / paragraph / table)."""

    kind: str  # "heading" | "paragraph" | "table"
    level: int  # 1/2/3 for headings, 0 otherwise
    text: str
    page_number: int | None = None


@dataclass
class LoadedDoc:
    source_type: str = "document"
    file_name: str = ""
    title: str = ""
    blocks: list[DocBlock] = field(default_factory=list)


def _strip_zero_width(text: str) -> str:
    return "".join(ch for ch in text if ch not in "\u200b\u200c\u200d\ufeff")


def _norm(text: str) -> str:
    return unicodedata.normalize("NFC", _strip_zero_width(text)).strip()


def _load_docx(path: Path) -> LoadedDoc:
    from docx import Document  # python-docx

    doc = Document(str(path))
    title = ""
    blocks: list[DocBlock] = []

    for para in doc.paragraphs:
        text = _norm(para.text)
        if not text:
            continue
        style = (para.style.name or "").lower() if para.style else ""
        if style.startswith("heading"):
            try:
                level = int(style.replace("heading", "").strip() or "1")
            except ValueError:
                level = 1
            if level == 1 and not title:
                title = text
            blocks.append(DocBlock(kind="heading", level=level, text=text))
        else:
            blocks.append(DocBlock(kind="paragraph", level=0, text=text))

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [_norm(c.text) for c in row.cells]
            rows.append(cells)
        if not rows:
            continue
        md = _rows_to_markdown(rows)
        blocks.append(DocBlock(kind="table", level=0, text=md))

    return LoadedDoc(file_name=path.name, title=title or path.stem, blocks=blocks)


def _rows_to_markdown(rows: list[list[str]]) -> str:
    if not rows:
        return ""
    header = rows[0]
    body = rows[1:]
    lines = ["| " + " | ".join(header) + " |"]
    lines.append("| " + " | ".join("---" for _ in header) + " |")
    for row in body:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _load_pdf(path: Path) -> LoadedDoc:
    import pdfplumber

    blocks: list[DocBlock] = []
    per_page_lines: list[list[str]] = []
    with pdfplumber.open(str(path)) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            lines = [_norm(line) for line in text.splitlines() if line.strip()]
            per_page_lines.append(lines)
            for table in page.extract_tables() or []:
                rows = [[_norm(c or "") for c in row] for row in table]
                blocks.append(
                    DocBlock(
                        kind="table",
                        level=0,
                        text=_rows_to_markdown(rows),
                        page_number=page_num,
                    )
                )

    repeated = _detect_repeated_lines(per_page_lines)
    for page_num, lines in enumerate(per_page_lines, start=1):
        for line in lines:
            if line in repeated:
                continue
            # Heuristic: short ALL-CAPS or short line ending with ':' treated as heading
            is_heading = len(line) < 80 and (
                line.isupper() or line.endswith(":")
            )
            blocks.append(
                DocBlock(
                    kind="heading" if is_heading else "paragraph",
                    level=2 if is_heading else 0,
                    text=line,
                    page_number=page_num,
                )
            )

    title = next(
        (b.text for b in blocks if b.kind == "heading" and b.level <= 2),
        path.stem,
    )
    return LoadedDoc(file_name=path.name, title=title, blocks=blocks)


def _detect_repeated_lines(pages: list[list[str]]) -> set[str]:
    """Lines that appear on more than half of the pages = header/footer."""
    if len(pages) < 3:
        return set()
    counts: Counter = Counter()
    for lines in pages:
        for line in set(lines):
            counts[line] += 1
    threshold = len(pages) // 2 + 1
    return {line for line, c in counts.items() if c >= threshold}


def _load_txt(path: Path) -> LoadedDoc:
    text = path.read_text(encoding="utf-8", errors="replace")
    blocks = [
        DocBlock(kind="paragraph", level=0, text=_norm(p))
        for p in text.split("\n\n")
        if p.strip()
    ]
    return LoadedDoc(file_name=path.name, title=path.stem, blocks=blocks)


def load_document(path: Path) -> LoadedDoc | None:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return _load_docx(path)
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix == ".txt":
        return _load_txt(path)
    return None


def load_documents(root: Path) -> Iterator[LoadedDoc]:
    if not root.exists():
        return
    for path in sorted(root.rglob("*")):
        if path.is_file() and path.suffix.lower() in {".docx", ".pdf", ".txt"}:
            try:
                doc = load_document(path)
            except Exception as exc:
                print(f"[doc_loader] failed on {path.name}: {exc}")
                continue
            if doc and doc.blocks:
                yield doc
